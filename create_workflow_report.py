import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 워드 문서 생성
doc = docx.Document()

# 문서 여백 설정
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# 제목 스타일 설정
title_style = doc.styles['Title']
title_style.font.name = 'HY견고딕'
title_style.font.size = Pt(20)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)

# 제목 추가
title = doc.add_paragraph('에스에스전력 휴가관리시스템 워크플로우 보고서', style='Title')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 날짜 추가
date_paragraph = doc.add_paragraph()
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
date_run = date_paragraph.add_run('2025년 5월 4일')
date_run.font.name = 'HY견고딕'
date_run.font.size = Pt(10)

doc.add_paragraph()  # 빈 줄 추가

# 소개 섹션
doc.add_heading('1. 개요', level=1)
intro = doc.add_paragraph()
intro.add_run('본 보고서는 에스에스전력 휴가관리시스템의 주요 기능 및 워크플로우를 상세히 기술합니다. 이 시스템은 직원들의 휴가 신청 프로세스를 디지털화하고, 관리자의 승인 절차를 효율적으로 관리하며, 재직증명서 발급과 같은 부가 기능을 제공합니다.')

# 시스템 아키텍처 섹션
doc.add_heading('2. 시스템 아키텍처', level=1)
arch = doc.add_paragraph()
arch.add_run('에스에스전력 휴가관리시스템은 Python과 Flask 프레임워크를 기반으로 개발된 웹 애플리케이션입니다. 주요 구성요소는 다음과 같습니다:')

arch_points = [
    "백엔드: Flask 웹 프레임워크를 사용하여 서버 측 로직 구현",
    "데이터베이스: PostgreSQL을 사용한 관계형 데이터베이스 관리",
    "프론트엔드: HTML, CSS, JavaScript, Tailwind CSS를 활용한 반응형 웹 인터페이스",
    "문서 생성: python-docx, reportlab 라이브러리를 통한 문서 생성 및 PDF 변환",
    "인증 및 권한 관리: Flask-Login을 이용한 사용자 인증 및 접근 제어"
]

arch_list = doc.add_paragraph(style='List Bullet')
for point in arch_points:
    if arch_list.text:  # 이미 내용이 있다면 새로운 항목 추가
        arch_list = doc.add_paragraph(style='List Bullet')
    arch_list.add_run(point)

# 데이터베이스 구조
doc.add_heading('3. 데이터베이스 구조', level=1)
db_intro = doc.add_paragraph()
db_intro.add_run('시스템은 다음과 같은 주요 데이터베이스 모델을 사용합니다:')

db_tables = [
    {
        'name': 'User (사용자)',
        'desc': '직원 및 관리자 계정 정보를 저장합니다.',
        'fields': [
            "id: 사용자 고유 식별자",
            "username: 로그인 아이디",
            "email: 이메일 주소",
            "password_hash: 암호화된 비밀번호",
            "name: 사용자 실명",
            "resident_id_first: 주민번호 앞자리",
            "resident_id_last_digit: 주민번호 뒷자리 첫번째 숫자",
            "role: 역할 (직원/관리자)",
            "department: 부서",
            "position: 직급",
            "hire_date: 입사일"
        ]
    },
    {
        'name': 'VacationDays (휴가일수)',
        'desc': '직원별 연간 휴가 일수를 관리합니다.',
        'fields': [
            "id: 고유 식별자",
            "user_id: 사용자 외래 키",
            "year: 연도",
            "total_days: 총 휴가 일수",
            "used_days: 사용한 휴가 일수"
        ]
    },
    {
        'name': 'VacationRequest (휴가신청)',
        'desc': '직원들의 휴가 신청 정보를 저장합니다.',
        'fields': [
            "id: 고유 식별자",
            "user_id: 신청자 외래 키",
            "start_date: 휴가 시작일",
            "end_date: 휴가 종료일",
            "days: 휴가 일수",
            "reason: 휴가 사유",
            "status: 상태 (대기중/승인됨/반려됨)",
            "type: 휴가 종류 (연차/반차/특별휴가)",
            "approved_by: 승인자 외래 키",
            "approval_date: 승인/반려 날짜",
            "comments: 관리자 코멘트",
            "created_at: 신청 날짜"
        ]
    },
    {
        'name': 'EmploymentCertificate (재직증명서)',
        'desc': '재직증명서 신청 및 발급 정보를 관리합니다.',
        'fields': [
            "id: 고유 식별자",
            "user_id: 신청자 외래 키",
            "purpose: 사용 목적",
            "issued_date: 발급일",
            "status: 상태 (대기중/발급완료/반려됨)",
            "approved_by: 승인자 외래 키",
            "approval_date: 승인/반려 날짜",
            "comments: 관리자 코멘트",
            "created_at: 신청 날짜"
        ]
    },
    {
        'name': 'Holiday (공휴일)',
        'desc': '공휴일 정보를 저장하여 휴가 일수 계산에 활용합니다.',
        'fields': [
            "id: 고유 식별자",
            "date: 공휴일 날짜",
            "name: 공휴일 이름"
        ]
    },
    {
        'name': 'CompanyInfo (회사정보)',
        'desc': '회사 정보를 저장하여, 재직증명서 등에 활용합니다.',
        'fields': [
            "id: 고유 식별자",
            "name: 회사명",
            "ceo_name: 대표자명",
            "registration_number: 사업자등록번호",
            "address: 회사 주소",
            "phone: 전화번호",
            "fax: 팩스번호",
            "website: 웹사이트",
            "stamp_image: 직인 이미지 (base64)"
        ]
    }
]

for table in db_tables:
    table_heading = doc.add_heading(level=2)
    table_heading.add_run(table['name'])
    
    desc = doc.add_paragraph()
    desc.add_run(table['desc'])
    
    for field in table['fields']:
        field_para = doc.add_paragraph(style='List Bullet')
        field_para.paragraph_format.left_indent = Pt(20)
        field_para.add_run(field)
    
    doc.add_paragraph()  # 빈 줄 추가

# 주요 기능 및 워크플로우
doc.add_heading('4. 주요 기능 및 워크플로우', level=1)

workflows = [
    {
        'name': '4.1 사용자 인증 워크플로우',
        'steps': [
            "사용자가 로그인 페이지 접속",
            "아이디와 비밀번호 입력",
            "서버에서 사용자 정보 검증",
            "역할에 따라 해당 대시보드로 리디렉션 (직원/관리자)",
            "로그인 실패 시 오류 메시지 표시"
        ]
    },
    {
        'name': '4.2 휴가 신청 워크플로우',
        'steps': [
            "직원이 휴가 신청 페이지 접속",
            "휴가 유형, 시작일, 종료일, 사유 입력",
            "휴가 일수 자동 계산 (주말, 공휴일 제외)",
            "입력 정보 검증 (남은 휴가 일수 확인, 중복 신청 확인)",
            "휴가 신청 정보 저장",
            "대기 상태로 관리자에게 통보",
            "직원 대시보드에 신청 내역 표시"
        ]
    },
    {
        'name': '4.3 휴가 승인 워크플로우',
        'steps': [
            "관리자가 휴가 관리 페이지 접속",
            "대기 중인 휴가 신청 목록 확인",
            "개별 휴가 신청 상세 내용 확인",
            "승인 또는 반려 결정 및 코멘트 입력",
            "승인 시 직원의 사용 휴가 일수 자동 업데이트",
            "처리 결과가 직원의 휴가 내역에 반영"
        ]
    },
    {
        'name': '4.4 재직증명서 신청 워크플로우',
        'steps': [
            "직원이 재직증명서 신청 페이지 접속",
            "사용 목적 입력",
            "신청 정보 저장",
            "대기 상태로 관리자에게 통보",
            "직원 대시보드에 신청 내역 표시"
        ]
    },
    {
        'name': '4.5 재직증명서 발급 워크플로우',
        'steps': [
            "관리자가 재직증명서 관리 페이지 접속",
            "대기 중인 재직증명서 신청 목록 확인",
            "개별 신청 상세 내용 확인",
            "승인 또는 반려 결정 및 코멘트 입력",
            "승인 시 재직증명서 자동 생성",
            "생성된 재직증명서는 직원이 다운로드 가능",
            "재직증명서에는 바코드와 발급일이 자동 포함"
        ]
    },
    {
        'name': '4.6 직원 관리 워크플로우',
        'steps': [
            "관리자가 직원 관리 페이지 접속",
            "등록된 직원 목록 확인",
            "개별 직원 정보 확인 및 수정",
            "직원 삭제 (관련 휴가, 재직증명서 기록 함께 삭제)",
            "엑셀 파일을 통한 직원 대량 등록",
            "직원별 연간 휴가 일수 설정",
            "직원 입사일 설정"
        ]
    },
    {
        'name': '4.7 회사 정보 관리 워크플로우',
        'steps': [
            "관리자가 회사 정보 관리 페이지 접속",
            "회사명, 대표자명, 사업자등록번호 등 정보 입력/수정",
            "직인 이미지 등록",
            "저장된 회사 정보는 재직증명서 생성 시 자동으로 적용"
        ]
    },
    {
        'name': '4.8 공휴일 관리 워크플로우',
        'steps': [
            "관리자가 공휴일 관리 페이지 접속",
            "연도별 공휴일 목록 확인",
            "새로운 공휴일 등록",
            "기존 공휴일 삭제",
            "등록된 공휴일은 휴가 일수 계산 시 자동으로 제외"
        ]
    }
]

for workflow in workflows:
    wf_heading = doc.add_heading(level=2)
    wf_heading.add_run(workflow['name'])
    
    for i, step in enumerate(workflow['steps'], 1):
        step_para = doc.add_paragraph()
        step_para.paragraph_format.left_indent = Pt(20)
        step_para.add_run(f"{i}. {step}")
    
    doc.add_paragraph()  # 빈 줄 추가

# 역할 및 권한
doc.add_heading('5. 역할 및 권한', level=1)

roles = [
    {
        'name': '5.1 직원 역할',
        'permissions': [
            "휴가 신청 및 취소",
            "본인의 휴가 내역 조회",
            "휴가 일수 확인",
            "재직증명서 신청 및 다운로드",
            "본인의 정보 확인"
        ]
    },
    {
        'name': '5.2 관리자 역할',
        'permissions': [
            "모든 직원 권한 포함",
            "직원 관리 (조회, 등록, 수정, 삭제)",
            "휴가 승인/반려",
            "재직증명서 발급 승인/반려",
            "공휴일 관리",
            "회사 정보 관리",
            "휴가 통계 확인 및 내보내기"
        ]
    }
]

for role in roles:
    role_heading = doc.add_heading(level=2)
    role_heading.add_run(role['name'])
    
    for perm in role['permissions']:
        perm_para = doc.add_paragraph(style='List Bullet')
        perm_para.add_run(perm)
    
    doc.add_paragraph()  # 빈 줄 추가

# 기술 스택
doc.add_heading('6. 기술 스택', level=1)

tech_stack = [
    {
        'name': '6.1 백엔드',
        'items': [
            "Python 3.11",
            "Flask 2.2.3 - 웹 프레임워크",
            "Flask-Login 0.6.2 - 사용자 인증",
            "Flask-SQLAlchemy 3.0.3 - ORM",
            "Flask-WTF 1.1.1 - 폼 처리",
            "SQLAlchemy 2.0.20 - 데이터베이스 ORM",
            "Werkzeug 2.2.3 - WSGI 유틸리티",
            "Gunicorn 21.2.0 - WSGI HTTP 서버"
        ]
    },
    {
        'name': '6.2 데이터베이스',
        'items': [
            "PostgreSQL - 관계형 데이터베이스",
            "psycopg2-binary 2.9.9 - PostgreSQL 어댑터"
        ]
    },
    {
        'name': '6.3 프론트엔드',
        'items': [
            "HTML5",
            "CSS3",
            "JavaScript",
            "Tailwind CSS - 반응형 UI 프레임워크"
        ]
    },
    {
        'name': '6.4 문서 생성',
        'items': [
            "python-docx 0.8.11 - 워드 문서 생성",
            "reportlab 4.0.4 - PDF 생성",
            "Pillow 10.0.0 - 이미지 처리"
        ]
    },
    {
        'name': '6.5 배포',
        'items': [
            "Render.com - 클라우드 호스팅",
            "Git - 버전 관리"
        ]
    }
]

for tech in tech_stack:
    tech_heading = doc.add_heading(level=2)
    tech_heading.add_run(tech['name'])
    
    for item in tech['items']:
        item_para = doc.add_paragraph(style='List Bullet')
        item_para.add_run(item)
    
    doc.add_paragraph()  # 빈 줄 추가

# 보안 및 데이터 보호
doc.add_heading('7. 보안 및 데이터 보호', level=1)

security_points = [
    "비밀번호 암호화: 사용자 비밀번호는 Werkzeug의 보안 해싱 기능을 사용하여 저장",
    "역할 기반 접근 제어: 사용자 역할에 따른 기능 및 데이터 접근 제한",
    "주민등록번호 보호: 주민번호는 앞 6자리와 뒷자리 첫 번째 숫자만 저장",
    "세션 관리: Flask-Login을 통한 안전한 사용자 세션 관리",
    "CSRF 보호: Flask-WTF의 CSRF 토큰을 사용한 크로스 사이트 요청 위조 방지",
    "입력값 검증: 모든 사용자 입력에 대한 서버 측 유효성 검사 수행"
]

for point in security_points:
    sec_para = doc.add_paragraph(style='List Bullet')
    sec_para.add_run(point)

# 결론
doc.add_heading('8. 결론', level=1)
conclusion = doc.add_paragraph()
conclusion.add_run('에스에스전력 휴가관리시스템은 직원들의 휴가 및 재직증명서 관리를 효율적으로 처리하기 위한 종합적인 솔루션입니다. Flask 기반의 안정적인 백엔드와 직관적인 사용자 인터페이스를 통해 관리자와 직원 모두에게 편리한 경험을 제공합니다. 한국 공휴일 자동 계산, 재직증명서 자동 생성, 직원 대량 등록 등의 기능을 통해 인사 업무 효율성을 크게 향상시키며, 부서별, 직급별 관리 체계를 통해 기업 내 인사 관리를 체계화합니다.')

# 저작권 문구 추가
footer_paragraph = doc.add_paragraph()
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_paragraph.add_run('© 2025 주식회사 에스에스전력 제작')
footer_run.font.name = 'HY견고딕'
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(100, 100, 100)

# 문서 저장
doc.save('에스에스전력_휴가관리시스템_워크플로우_보고서.docx')

print("워크플로우 보고서가 성공적으로 생성되었습니다.")