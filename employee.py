from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify, make_response
from flask_login import login_required, current_user
from app import db
from models import VacationDays, VacationRequest, VacationStatus, EmploymentCertificate, CertificateStatus, CompanyInfo
from forms import VacationRequestForm, EmploymentCertificateRequestForm, VacationSearchForm
from datetime import datetime
from utils import get_vacation_days_count, check_overlapping_vacation
import tempfile
import os
import urllib.parse
import io
import pandas as pd
# WeasyPrint는 시스템 종속성 문제가 있어 사용하지 않습니다
# 필요한 경우 reportlab을 사용합니다
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw, ImageFont
import base64
import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
# qrcode 모듈은 더 이상 사용하지 않습니다
import PIL
from PIL import Image, ImageDraw, ImageFont

employee_bp = Blueprint('employee', __name__)

@employee_bp.route('/dashboard')
@login_required
def dashboard():
    """직원 대시보드"""
    # 현재 연도의 휴가 정보 가져오기
    current_year = datetime.now().year
    vacation_days = VacationDays.query.filter_by(
        user_id=current_user.id,
        year=current_year
    ).first()
    
    # 없으면 생성
    if not vacation_days:
        vacation_days = VacationDays(
            user_id=current_user.id,
            year=current_year,
            total_days=15,  # 기본값
            used_days=0
        )
        db.session.add(vacation_days)
        db.session.commit()
    
    # 최근 휴가 신청 내역 (5개)
    recent_requests = VacationRequest.query.filter_by(
        user_id=current_user.id
    ).order_by(VacationRequest.created_at.desc()).limit(5).all()
    
    # 대기중인 휴가 신청 수
    pending_count = VacationRequest.query.filter_by(
        user_id=current_user.id,
        status=VacationStatus.PENDING
    ).count()
    
    # 승인된 휴가 신청 수
    approved_count = VacationRequest.query.filter_by(
        user_id=current_user.id,
        status=VacationStatus.APPROVED
    ).count()
    
    # 추가 통계 데이터
    remaining_vacation_days = vacation_days.remaining_days()
    total_vacation_days = vacation_days.total_days
    
    # 대기중인 모든 신청 (휴가 + 증명서)
    pending_requests = (VacationRequest.query.filter_by(user_id=current_user.id, status=VacationStatus.PENDING).count() +
                       EmploymentCertificate.query.filter_by(user_id=current_user.id, status=CertificateStatus.PENDING).count())
    
    # 최근 휴가 신청 내역 (5개)
    recent_my_vacations = VacationRequest.query.filter_by(
        user_id=current_user.id
    ).order_by(VacationRequest.created_at.desc()).limit(5).all()
    
    # 이번 달 예정된 이벤트 (휴가, 공휴일)
    from models import Holiday
    current_month = datetime.now().month
    current_year = datetime.now().year
    
    upcoming_events = []
    
    # 승인된 휴가 일정
    approved_vacations = VacationRequest.query.filter_by(
        user_id=current_user.id,
        status=VacationStatus.APPROVED
    ).filter(
        VacationRequest.start_date >= datetime.now().date()
    ).order_by(VacationRequest.start_date).limit(10).all()
    
    for vacation in approved_vacations:
        upcoming_events.append({
            'date': vacation.start_date,
            'type': 'vacation',
            'description': f'{vacation.type} ({vacation.days}일)'
        })
    
    # 공휴일
    holidays = Holiday.query.filter(
        Holiday.date >= datetime.now().date()
    ).order_by(Holiday.date).limit(10).all()
    
    for holiday in holidays:
        upcoming_events.append({
            'date': holiday.date,
            'type': 'holiday',
            'description': holiday.name
        })
    
    # 날짜순 정렬
    upcoming_events.sort(key=lambda x: x['date'])
    upcoming_events = upcoming_events[:10]  # 최대 10개

    return render_template(
        'employee/dashboard_gov.html',
        remaining_vacation_days=remaining_vacation_days,
        total_vacation_days=total_vacation_days,
        pending_requests=pending_requests,
        recent_my_vacations=recent_my_vacations,
        upcoming_events=upcoming_events,
        current_year=current_year
    )

@employee_bp.route('/request-vacation', methods=['GET', 'POST'])
@login_required
def request_vacation():
    """휴가 신청 페이지"""
    form = VacationRequestForm()
    
    if form.validate_on_submit():
        # 현재 연도의 휴가 정보 가져오기
        year = form.start_date.data.year
        vacation_days = VacationDays.query.filter_by(
            user_id=current_user.id,
            year=year
        ).first()
        
        # 없으면 생성
        if not vacation_days:
            vacation_days = VacationDays(
                user_id=current_user.id,
                year=year,
                total_days=15,  # 기본값
                used_days=0
            )
            db.session.add(vacation_days)
            db.session.commit()
        
        # 휴가 일수 계산
        type_days = {
            '연차': float(form.days.data),
            '반차(오전)': 0.5,
            '반차(오후)': 0.5,
            '특별휴가': float(form.days.data)
        }
        days = type_days.get(form.type.data, float(form.days.data))
        
        # 남은 휴가 일수 확인 (특별휴가는 연차 차감 없음)
        if form.type.data != '특별휴가' and days > vacation_days.remaining_days():
            flash('남은 휴가 일수가 부족합니다.', 'danger')
            return render_template('employee/request_vacation.html', form=form)
        
        # 같은 기간에 이미 신청한 휴가가 있는지 확인
        if check_overlapping_vacation(current_user.id, form.start_date.data, form.end_date.data):
            flash('이미 해당 기간에 신청한 휴가가 있습니다.', 'danger')
            return render_template('employee/request_vacation.html', form=form)
        
        # 휴가 신청 저장
        vacation_request = VacationRequest(
            user_id=current_user.id,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            days=days,
            reason=form.reason.data,
            type=form.type.data,
            status=VacationStatus.PENDING
        )
        
        db.session.add(vacation_request)
        db.session.commit()
        
        flash('휴가 신청이 완료되었습니다.', 'success')
        return redirect(url_for('employee.my_vacations'))
    
    return render_template('employee/request_vacation.html', form=form)

@employee_bp.route('/my-vacations', methods=['GET', 'POST'])
@login_required
def my_vacations():
    """내 휴가 내역 페이지 (기간 검색 및 엑셀 출력 지원)"""
    form = VacationSearchForm()
    
    # 기본 쿼리
    query = VacationRequest.query.filter_by(user_id=current_user.id)
    
    # 기본 연도 설정
    current_year = datetime.now().year
    search_year = current_year
    
    # 폼 처리
    if form.validate_on_submit():
        # 엑셀 다운로드 요청
        if form.export.data:
            return export_my_vacation_data(form)
        
        # 검색 필터 적용
        if form.start_date.data:
            query = query.filter(VacationRequest.start_date >= form.start_date.data)
        
        if form.end_date.data:
            query = query.filter(VacationRequest.end_date <= form.end_date.data)
        
        if form.status.data != 'all':
            query = query.filter(VacationRequest.status == form.status.data)
        
        if form.year.data != 0:
            query = query.filter(db.extract('year', VacationRequest.start_date) == form.year.data)
            search_year = form.year.data
    else:
        # URL 파라미터로부터 필터 적용 (기존 호환성)
        year = request.args.get('year', current_year, type=int)
        status = request.args.get('status', 'all')
        
        if year != current_year:
            query = query.filter(db.extract('year', VacationRequest.start_date) == year)
            form.year.data = year
            search_year = year
        
        if status != 'all':
            query = query.filter(VacationRequest.status == status)
            form.status.data = status
    
    # 정렬 (최신순)
    vacation_requests = query.order_by(VacationRequest.created_at.desc()).all()
    
    # 연도별 휴가 정보
    vacation_days = VacationDays.query.filter_by(
        user_id=current_user.id,
        year=search_year
    ).first()
    
    if not vacation_days:
        vacation_days = VacationDays(
            user_id=current_user.id,
            year=search_year,
            total_days=15,  # 기본값
            used_days=0
        )
        db.session.add(vacation_days)
        db.session.commit()
    
    return render_template(
        'employee/my_vacations_gov.html',
        vacation_requests=vacation_requests,
        vacation_days=vacation_days,
        year=search_year,
        status_filter=form.status.data or 'all',
        search_form=form
    )

@employee_bp.route('/cancel-vacation/<int:request_id>', methods=['POST'])
@login_required
def cancel_vacation(request_id):
    """휴가 신청 취소"""
    vacation_request = VacationRequest.query.get_or_404(request_id)
    
    # 권한 확인
    if vacation_request.user_id != current_user.id:
        flash('권한이 없습니다.', 'danger')
        return redirect(url_for('employee.my_vacations'))
    
    # 대기 중인 신청만 취소 가능
    if vacation_request.status != VacationStatus.PENDING:
        flash('대기 중인 휴가 신청만 취소할 수 있습니다.', 'danger')
        return redirect(url_for('employee.my_vacations'))
    
    # 취소 처리
    db.session.delete(vacation_request)
    db.session.commit()
    
    flash('휴가 신청이 취소되었습니다.', 'success')
    return redirect(url_for('employee.my_vacations'))


def export_my_vacation_data(form):
    """개인 휴가 데이터 엑셀 다운로드"""
    # 쿼리 생성
    query = VacationRequest.query.filter_by(user_id=current_user.id)
    
    # 검색 조건 적용
    if form.start_date.data:
        query = query.filter(VacationRequest.start_date >= form.start_date.data)
    
    if form.end_date.data:
        query = query.filter(VacationRequest.end_date <= form.end_date.data)
    
    if form.status.data != 'all':
        query = query.filter(VacationRequest.status == form.status.data)
    
    if form.year.data != 0:
        query = query.filter(db.extract('year', VacationRequest.start_date) == form.year.data)
    
    # 정렬
    query = query.order_by(VacationRequest.created_at.desc())
    
    # 결과 가져오기
    results = query.all()
    
    # 엑셀 파일 생성 (pandas 이용)
    data = []
    for vacation_request in results:
        data.append({
            '휴가시작일': vacation_request.start_date.strftime('%Y-%m-%d'),
            '휴가종료일': vacation_request.end_date.strftime('%Y-%m-%d'),
            '휴가일수': vacation_request.days,
            '휴가유형': vacation_request.type,
            '휴가사유': vacation_request.reason or '',
            '상태': vacation_request.status,
            '신청일시': vacation_request.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            '승인일시': vacation_request.approval_date.strftime('%Y-%m-%d %H:%M:%S') if vacation_request.approval_date else ''
        })
    
    # DataFrame 생성
    df = pd.DataFrame(data)
    
    # 파일명 생성
    current_time = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'{current_user.name}_휴가현황_{current_time}.xlsx'
    
    # 임시 파일로 저장
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, engine='openpyxl')
        
        # 파일 읽기
        with open(tmp.name, 'rb') as f:
            excel_data = f.read()
        
        # 임시 파일 삭제
        os.unlink(tmp.name)
    
    # 응답 생성
    response = make_response(excel_data)
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

@employee_bp.route('/calculate-vacation-days', methods=['POST'])
@login_required
def calculate_vacation_days():
    """휴가 일수 계산 API"""
    data = request.json
    start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date()
    end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date()
    vacation_type = data.get('type', '연차')
    
    if vacation_type in ['반차(오전)', '반차(오후)']:
        days = 0.5
    else:
        days = get_vacation_days_count(start_date, end_date)
    
    return jsonify({'days': days})


@employee_bp.route('/request-certificate', methods=['GET', 'POST'])
@login_required
def request_certificate():
    """재직증명서 신청 페이지"""
    form = EmploymentCertificateRequestForm()
    
    if form.validate_on_submit():
        # 재직증명서 신청 저장
        certificate = EmploymentCertificate(
            user_id=current_user.id,
            purpose=form.purpose.data,
            status=CertificateStatus.PENDING
        )
        
        db.session.add(certificate)
        db.session.commit()
        
        flash('재직증명서 신청이 완료되었습니다.', 'success')
        return redirect(url_for('employee.my_certificates'))
    
    return render_template('employee/request_certificate.html', form=form)


@employee_bp.route('/my-certificates')
@login_required
def my_certificates():
    """내 재직증명서 신청 내역 페이지"""
    # 재직증명서 신청 내역 (최신순)
    certificates = EmploymentCertificate.query.filter_by(
        user_id=current_user.id
    ).order_by(EmploymentCertificate.created_at.desc()).all()
    
    return render_template('employee/my_certificates.html', certificates=certificates)


def text_to_image(text, font_size=24, width=None, text_color=(0, 0, 0), bg_color=(255, 255, 255)):
    """텍스트를 이미지로 변환하는 함수"""
    # 기본 폰트 (Arial)
    try:
        # 시스템에 설치된 폰트 사용
        font = ImageFont.truetype("Arial", font_size)
    except:
        # 기본 폰트 사용
        font = ImageFont.load_default()
    
    # 텍스트 크기 계산
    test_img = Image.new('RGB', (1, 1))
    test_draw = ImageDraw.Draw(test_img)
    text_width, text_height = test_draw.textsize(text, font=font) if hasattr(test_draw, 'textsize') else font.getbbox(text)[2:4]
    
    # 이미지 크기 설정
    if width:
        img_width = width
    else:
        img_width = text_width + 20  # 여백 추가
    img_height = text_height + 20  # 여백 추가
    
    # 이미지 생성
    img = Image.new('RGB', (img_width, img_height), color=bg_color)
    draw = ImageDraw.Draw(img)
    
    # 텍스트 추가 (가운데 정렬)
    x = (img_width - text_width) // 2
    y = (img_height - text_height) // 2
    draw.text((x, y), text, font=font, fill=text_color)
    
    # 이미지를 바이트로 변환
    img_byte = io.BytesIO()
    img.save(img_byte, format='PNG')
    img_byte.seek(0)
    
    return img_byte.getvalue()

def create_qrcode(data, size=200):
    """QR 코드 생성 함수 (간단한 격자무늬 이미지로 대체)"""
    # 이미지 크기와 색상 설정
    img_size = size
    background_color = (255, 255, 255)  # 흰색
    qrcode_color = (0, 0, 0)  # 검은색
    
    # 새 이미지 생성
    img = Image.new('RGB', (img_size, img_size), color=background_color)
    draw = ImageDraw.Draw(img)
    
    # 격자 크기 계산
    grid_size = img_size // 25
    
    # 데이터를 이용한 패턴 생성
    data_hash = sum([ord(c) for c in data])
    
    # 격자 그리기
    for x in range(0, img_size, grid_size):
        for y in range(0, img_size, grid_size):
            # 데이터 기반 패턴 생성
            if ((x // grid_size) * 5 + (y // grid_size) + data_hash) % 3 == 0:
                draw.rectangle(
                    (x, y, x + grid_size - 1, y + grid_size - 1),
                    fill=qrcode_color
                )
    
    # 테두리 추가
    draw.rectangle((0, 0, img_size-1, img_size-1), outline=qrcode_color, width=2)
    
    # QR 코드 영역 표시 (코너 사각형)
    corner_size = grid_size * 3
    
    # 왼쪽 상단
    draw.rectangle((0, 0, corner_size, corner_size), fill=qrcode_color)
    draw.rectangle((grid_size, grid_size, corner_size-grid_size, corner_size-grid_size), fill=background_color)
    
    # 오른쪽 상단
    draw.rectangle((img_size-corner_size-1, 0, img_size-1, corner_size), fill=qrcode_color)
    draw.rectangle((img_size-corner_size+grid_size-1, grid_size, img_size-grid_size-1, corner_size-grid_size), fill=background_color)
    
    # 왼쪽 하단
    draw.rectangle((0, img_size-corner_size-1, corner_size, img_size-1), fill=qrcode_color)
    draw.rectangle((grid_size, img_size-corner_size+grid_size-1, corner_size-grid_size, img_size-grid_size-1), fill=background_color)
    
    # 이미지를 바이트로 변환
    img_byte = io.BytesIO()
    img.save(img_byte, format='PNG')
    img_byte.seek(0)
    
    return img_byte

def create_barcode(data, width=400, height=100):
    """바코드 생성 함수 (가로형 형태)"""
    # 최소 높이 보장 (오류 방지)
    height = max(height, 30)
    
    # 이미지 설정
    img_width = width
    img_height = height
    background_color = (255, 255, 255)  # 흰색
    barcode_color = (0, 0, 0)  # 검은색
    
    # 새 이미지 생성
    img = Image.new('RGB', (img_width, img_height), color=background_color)
    draw = ImageDraw.Draw(img)
    
    # 바코드 패턴 설정
    bar_width = 3  # 바코드 선 너비
    gap_width = 2  # 간격
    max_bars = 35  # 바 개수
    position = 20  # 시작 위치
    bar_height = max(img_height - 25, 10)  # 바코드 높이 (최소값 보장)
    
    # 데이터 해시 값으로 일관된 패턴 생성
    data_hash = sum([ord(c) for c in data])
    
    # 바코드 생성
    for i in range(max_bars):
        # 데이터 기반 패턴 (높이 변화 축소)
        height_variation = min(((i * 5 + data_hash) % 3), 1)  # 높이 변화 최소화
        actual_bar_height = bar_height - height_variation

        # 바 그리기 - 특정 패턴으로 바 생성
        if ((i * 11 + data_hash) % 5) != 0:
            draw.rectangle(
                (position, 5, position + bar_width, 5 + actual_bar_height),
                fill=barcode_color
            )

        # 가끔 더 굵은 바 추가
        if i % 8 == 0:
            draw.rectangle(
                (position + bar_width + gap_width, 5, 
                 position + bar_width * 2 + gap_width, 5 + actual_bar_height),
                fill=barcode_color
            )
            position += bar_width + gap_width
            
        position += bar_width + gap_width
    
    # 바코드 번호 추가
    try:
        font = ImageFont.truetype("Arial", 10)  # 폰트 크기 줄임
    except:
        font = ImageFont.load_default()
    
    # 바코드 번호 텍스트 추가 (바닥에 가깝게 배치)
    text_width = draw.textbbox((0, 0), data, font=font)[2]
    text_x = (img_width - text_width) // 2
    # 텍스트 위치 조정 - 이미지 바닥에 더 가깝게
    draw.text((text_x, img_height - 15), data, font=font, fill=barcode_color)
    
    # 테두리 추가
    draw.rectangle(
        (10, 2, img_width - 10, img_height - 17),
        outline=barcode_color,
        width=1
    )
    
    # 이미지를 바이트로 변환
    img_byte = io.BytesIO()
    img.save(img_byte, format='PNG', dpi=(300, 300))
    img_byte.seek(0)
    
    return img_byte


def create_docx_certificate(certificate, current_user, company_info):
    """워드 파일로 재직증명서 생성 - 이미지와 정확히 동일한 형식"""
    company_name = company_info.name if company_info else '주식회사 에스에스전력'
    ceo_name = company_info.ceo_name if company_info else '김세인'
    
    today = datetime.now().date()
    today_str = f"{today.year}년 {today.month}월 {today.day}일"
    
    hire_date_str = ""
    if current_user.hire_date:
        hire_date_str = current_user.hire_date.strftime('%Y년 %m월 %d일')
    else:
        hire_date_str = "2024년 12월 20일"  # 기본값 설정
    
    # 워드 문서 생성
    doc = docx.Document()
    
    # A4 사이즈 설정 (단위: cm)
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # 여백을 이미지와 유사하게 설정
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'HY견고딕'
    style.font.size = Pt(10)
    
    # 발급일을 문서 상단 우측 끝으로 배치
    issue_date_p = doc.add_paragraph()
    issue_date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    issue_date_p.space_after = Pt(20)  # 아래쪽 공백 추가
    issue_date_run = issue_date_p.add_run(f'발급일: {today_str}')
    issue_date_run.font.name = 'HY견고딕'
    issue_date_run.font.size = Pt(10)
    
    # 제목 추가 - 중앙 정렬
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('재직증명서')
    title_run.font.name = 'HY견고딕'
    title_run.font.size = Pt(20)  # 폰트 크기 20으로 변경
    title_run.font.bold = True
    title.space_after = Pt(12)  # 제목 아래 약간의 여백
    
    # 가로선 추가 (테이블 방식으로 구현)
    line_table = doc.add_table(rows=1, cols=1)
    line_table.style = 'Table Grid'
    line_table.autofit = False
    line_table.width = Cm(16)
    line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 테이블 높이를 아주 작게 설정해서 선으로 보이게 함
    line_table.rows[0].height = Cm(0.05)
    line_table.rows[0].height_rule = 2  # WD_ROW_HEIGHT.EXACTLY = 2
    # 여백 추가
    p_after_line = doc.add_paragraph()
    p_after_line.space_after = Pt(12)
    
    # 표 생성
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    table.width = Cm(16)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 표 내용 채우기 (이미지와 정확히 동일하게)
    # 1행: 성명, 주민등록번호
    row = table.rows[0]
    row.cells[0].text = '성명'
    row.cells[1].text = current_user.name if current_user.name else '김영희'  # 기본값
    row.cells[2].text = '주민등록번호'
    
    # 주민번호 표시 (실제 주민번호가 있으면 표시, 없으면 기본값)
    if current_user.resident_id_first and current_user.resident_id_last_digit:
        resident_id = f"{current_user.resident_id_first}-{current_user.resident_id_last_digit}******"
    else:
        resident_id = "******-*******"  # 기본값
    row.cells[3].text = resident_id
    
    # 2행: 소속, 직위
    row = table.rows[1]
    row.cells[0].text = '소속'
    row.cells[1].text = current_user.department or '개발팀'
    row.cells[2].text = '직위'
    row.cells[3].text = current_user.position or '과장'
    
    # 3행: 재직기간
    row = table.rows[2]
    row.cells[0].text = '재직기간'
    cell = row.cells[1]
    cell.merge(row.cells[2])
    cell.merge(row.cells[3])
    cell.text = f"{hire_date_str} ~ 현재"
    
    # 4행: 용도
    row = table.rows[3]
    row.cells[0].text = '용도'
    cell = row.cells[1]
    cell.merge(row.cells[2])
    cell.merge(row.cells[3])
    cell.text = certificate.purpose if certificate and certificate.purpose else '개인'
    
    # 표 셀 스타일 설정
    for row in table.rows:
        row.height = Cm(0.9)  # 행 높이 설정
        row.height_rule = 2  # WD_ROW_HEIGHT.EXACTLY = 2
        
        for cell in row.cells:
            cell._element.tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.space_before = Pt(0)
                paragraph.space_after = Pt(0)
                
                for run in paragraph.runs:
                    run.font.name = 'HY견고딕'
                    run.font.size = Pt(12)
    
    # 표 아래 추가 공백 (네 칸으로 증가)
    for i in range(4):
        table_space = doc.add_paragraph()
        table_space.space_before = Pt(15)
        table_space.space_after = Pt(0)
    
    # 증명 문구 (표 아래 네 칸 이후에 배치)
    p_confirm = doc.add_paragraph()
    p_confirm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_confirm.space_before = Pt(0)
    p_confirm.space_after = Pt(25)
    confirm_run = p_confirm.add_run("상기인은 위와 같이 재직하고 있음을 증명합니다.")
    confirm_run.font.name = 'HY견고딕'
    confirm_run.font.size = Pt(13)  # 폰트 크기 13으로 변경
    
    # 중앙 여백 (이미지처럼 더 많은 간격 추가)
    for i in range(2):
        empty_p = doc.add_paragraph()
        empty_p.space_before = Pt(10)
        empty_p.space_after = Pt(0)
    
    # 날짜를 회사명 위로 배치 (중앙 정렬)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.space_before = Pt(10)
    date_p.space_after = Pt(10)
    date_run = date_p.add_run(today_str)
    date_run.font.name = 'HY견고딕'
    date_run.font.size = Pt(14)  # 폰트 크기 14로 변경
    
    # 날짜와 회사이름 사이 추가 여백
    date_company_space = doc.add_paragraph()
    date_company_space.space_before = Pt(15)
    
    # 회사명
    company_p = doc.add_paragraph()
    company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_p.space_before = Pt(0)
    company_p.space_after = Pt(0)
    company_run = company_p.add_run(company_name)
    company_run.font.name = 'HY견고딕'
    company_run.font.size = Pt(15)  # 폰트 크기 15로 변경
    company_run.font.bold = True
    
    # 회사 이름과 대표이사 이름을 바로 붙여서 배치 (여백 제거)
    
    # 대표이사 및 직인 생략
    ceo_p = doc.add_paragraph()
    ceo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ceo_p.space_before = Pt(0)
    ceo_p.space_after = Pt(0)
    ceo_run = ceo_p.add_run(f"대표이사 김세인")  # 이름 수정
    ceo_run.font.name = 'HY견고딕'
    ceo_run.font.size = Pt(15)  # 폰트 크기 15로 변경
    ceo_run.font.bold = True
    
    ceo_p.add_run(" ")  # 간격 추가
    
    seal_run = ceo_p.add_run("(직인 생략)")
    seal_run.font.name = 'HY견고딕'
    seal_run.font.size = Pt(12)  # 폰트 크기 12로 변경
    
    # 추가 공백 삽입 (원본 확인 관련 부분을 더 아래로 내리기)
    spacer_p = doc.add_paragraph()
    spacer_p.space_before = Pt(10)
    spacer_p.space_after = Pt(10)
    
    # 문서 확인번호 생성 (이미지와 동일한 형식)
    doc_verification_code = f"CERT-2-2-{datetime.now().strftime('%Y%m%d')}"
    
    # 원본 확인 안내문
    verify_note_p = doc.add_paragraph()
    verify_note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verify_note_p.space_before = Pt(0)
    verify_note_p.space_after = Pt(6)
    verify_note_run = verify_note_p.add_run("※ 아래 바코드로 문서의 진위여부를 확인하실 수 있습니다.")
    verify_note_run.font.name = 'HY견고딕'
    verify_note_run.font.size = Pt(8)
    verify_note_run.font.bold = True
    
    # 가로형 바코드 생성 및 삽입
    try:
        # 표 너비에 맞춘 가로형 바코드 생성 (세로 크기 60으로 설정)
        barcode_img_io = create_barcode(doc_verification_code, width=500, height=60)
        
        barcode_p = doc.add_paragraph()
        barcode_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        barcode_p.space_before = Pt(3)
        barcode_p.space_after = Pt(6)
        
        barcode_run = barcode_p.add_run()
        barcode_run.add_picture(barcode_img_io, width=Cm(16), height=Cm(1.5))  # 세로 크기 1.5cm로 설정
    except Exception as e:
        print(f"바코드 생성 오류: {str(e)}")
    
    # 문서확인번호
    code_p = doc.add_paragraph()
    code_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_p.space_before = Pt(2)
    code_p.space_after = Pt(0)
    code_run = code_p.add_run(f"문서확인번호: {doc_verification_code}")
    code_run.font.name = 'HY견고딕'
    code_run.font.size = Pt(8)
    
    # 문서확인 사이트
    guide_p = doc.add_paragraph()
    guide_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_p.space_before = Pt(1)
    guide_p.space_after = Pt(0)
    guide_run = guide_p.add_run(f"문서확인 사이트: {company_info.website if company_info and company_info.website else 'https://ss-electric.co.kr'}")
    guide_run.font.name = 'HY견고딕'
    guide_run.font.size = Pt(8)
    
    # 문서 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


@employee_bp.route('/download-certificate/<int:certificate_id>')
@login_required
def download_certificate(certificate_id):
    """재직증명서 다운로드"""
    certificate = EmploymentCertificate.query.get_or_404(certificate_id)
    
    # 권한 확인
    if certificate.user_id != current_user.id:
        flash('권한이 없습니다.', 'danger')
        return redirect(url_for('employee.my_certificates'))
    
    # 발급완료 상태인지 확인
    if certificate.status != CertificateStatus.ISSUED:
        flash('아직 발급되지 않은 재직증명서입니다.', 'warning')
        return redirect(url_for('employee.my_certificates'))
    
    # 회사 정보 가져오기
    company_info = CompanyInfo.query.first()
    company_name = company_info.name if company_info else '주식회사 에스에스전력'
    ceo_name = company_info.ceo_name if company_info else '대표이사'
    
    try:
        # 워드 문서 생성
        buffer = create_docx_certificate(certificate, current_user, company_info)
        
        # 응답 생성
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # 파일명 생성 및 인코딩
        filename = f'재직증명서_{current_user.name}_{datetime.now().strftime("%Y%m%d")}.docx'
        encoded_filename = urllib.parse.quote(filename)
        response.headers['Content-Disposition'] = f'attachment; filename={encoded_filename}'
        
        return response
    except Exception as e:
        # 오류 발생 시 로그 출력 및 처리
        print(f"파일 생성 오류: {str(e)}")
        flash('파일 생성 중 오류가 발생했습니다.', 'danger')
        return redirect(url_for('employee.my_certificates'))


@employee_bp.route('/cancel-certificate/<int:certificate_id>', methods=['POST'])
@login_required
def cancel_certificate(certificate_id):
    """재직증명서 신청 취소"""
    certificate = EmploymentCertificate.query.get_or_404(certificate_id)
    
    # 권한 확인
    if certificate.user_id != current_user.id:
        flash('권한이 없습니다.', 'danger')
        return redirect(url_for('employee.my_certificates'))
    
    # 대기 중인 신청만 취소 가능
    if certificate.status != CertificateStatus.PENDING:
        flash('대기 중인 재직증명서 신청만 취소할 수 있습니다.', 'danger')
        return redirect(url_for('employee.my_certificates'))
    
    # 취소 처리
    db.session.delete(certificate)
    db.session.commit()
    
    flash('재직증명서 신청이 취소되었습니다.', 'success')
    return redirect(url_for('employee.my_certificates'))
