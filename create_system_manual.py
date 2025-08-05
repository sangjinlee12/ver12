#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
에스에스전력 휴가관리시스템 설명서 생성 스크립트
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

def add_heading_numbered(doc, text, level=1):
    """번호가 있는 제목 추가"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_style(doc, text, style_name=None):
    """스타일이 적용된 단락 추가"""
    paragraph = doc.add_paragraph(text)
    if style_name:
        paragraph.style = style_name
    return paragraph

def create_table_with_style(doc, rows, cols, data=None):
    """스타일이 적용된 표 생성"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 헤더 스타일 적용
    if data:
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)
                if i == 0:  # 헤더 행
                    cell.paragraphs[0].runs[0].bold = True
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    return table

def create_system_manual():
    """시스템 설명서 생성"""
    doc = Document()
    
    # 문서 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    
    # 제목 페이지
    title = doc.add_heading('에스에스전력 휴가관리시스템', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('시스템 운영 매뉴얼', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 날짜 및 버전 정보
    doc.add_paragraph()
    version_para = doc.add_paragraph(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version_para = doc.add_paragraph('버전: 1.0')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 페이지 나누기
    doc.add_page_break()
    
    # 목차 (수동 작성)
    add_heading_numbered(doc, '목차', level=1)
    toc_data = [
        '1. 시스템 개요',
        '2. 시스템 접속 및 로그인',
        '3. 사용자 계정 관리',
        '4. 직원 기능',
        '   4.1 대시보드',
        '   4.2 휴가 신청',
        '   4.3 나의 휴가현황',
        '   4.4 재직증명서 신청',
        '5. 관리자 기능',
        '   5.1 관리자 대시보드',
        '   5.2 직원 관리',
        '   5.3 휴가 관리',
        '   5.4 재직증명서 관리',
        '   5.5 공휴일 관리',
        '6. 고급 기능',
        '   6.1 기간 검색',
        '   6.2 엑셀 출력',
        '   6.3 계정 복구',
        '7. 시스템 관리',
        '8. 문제해결 가이드'
    ]
    
    for item in toc_data:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. 시스템 개요
    add_heading_numbered(doc, '1. 시스템 개요', level=1)
    
    doc.add_paragraph('에스에스전력 휴가관리시스템은 회사 직원들의 휴가 신청, 승인, 관리를 위한 웹 기반 시스템입니다.')
    
    add_heading_numbered(doc, '1.1 주요 기능', level=2)
    features = [
        '직원 휴가 신청 및 관리',
        '관리자 휴가 승인/반려 처리',
        '재직증명서 발급',
        '공휴일 자동 관리',
        '기간별 휴가 현황 검색',
        '엑셀 파일 다운로드',
        '사용자 계정 복구 기능'
    ]
    
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    add_heading_numbered(doc, '1.2 시스템 구성', level=2)
    doc.add_paragraph('• 웹 브라우저 기반 접속')
    doc.add_paragraph('• 대한민국 정부 웹사이트 디자인 적용')
    doc.add_paragraph('• 모바일 반응형 디자인')
    doc.add_paragraph('• 안전한 데이터 보관 및 관리')
    
    # 2. 시스템 접속 및 로그인
    add_heading_numbered(doc, '2. 시스템 접속 및 로그인', level=1)
    
    add_heading_numbered(doc, '2.1 시스템 접속', level=2)
    doc.add_paragraph('1. 웹 브라우저에서 시스템 주소에 접속합니다.')
    doc.add_paragraph('2. 메인 페이지에서 "로그인" 버튼을 클릭합니다.')
    
    add_heading_numbered(doc, '2.2 로그인 방법', level=2)
    doc.add_paragraph('1. 아이디와 비밀번호를 입력합니다.')
    doc.add_paragraph('2. "로그인" 버튼을 클릭합니다.')
    doc.add_paragraph('3. 로그인 성공 시 사용자 권한에 따라 대시보드로 이동합니다.')
    
    add_heading_numbered(doc, '2.3 계정 복구', level=2)
    doc.add_paragraph('로그인 페이지에서 아이디나 비밀번호를 잊어버린 경우:')
    doc.add_paragraph('• "아이디 찾기": 이름과 이메일을 입력하여 아이디 확인')
    doc.add_paragraph('• "비밀번호 찾기": 아이디와 이메일을 입력하여 임시 비밀번호 발급')
    
    # 3. 사용자 계정 관리
    add_heading_numbered(doc, '3. 사용자 계정 관리', level=1)
    
    add_heading_numbered(doc, '3.1 회원가입', level=2)
    doc.add_paragraph('신규 직원은 다음 절차로 계정을 생성합니다:')
    doc.add_paragraph('1. 메인 페이지에서 "회원가입" 클릭')
    doc.add_paragraph('2. 필수 정보 입력:')
    doc.add_paragraph('   • 아이디 (4-20자)')
    doc.add_paragraph('   • 이메일')
    doc.add_paragraph('   • 이름')
    doc.add_paragraph('   • 주민번호 앞자리 및 뒷자리 첫째 자리')
    doc.add_paragraph('   • 비밀번호 (6자 이상)')
    doc.add_paragraph('   • 부서 및 직급 선택')
    doc.add_paragraph('3. "회원가입" 버튼 클릭')
    
    add_heading_numbered(doc, '3.2 부서 및 직급', level=2)
    
    # 부서 표
    doc.add_paragraph('등록 가능한 부서:')
    dept_data = [
        ['부서명', '설명'],
        ['공사팀', '시설 공사 업무'],
        ['공무부', '공무 관련 업무'],
        ['경리부', '회계 및 재무 업무'],
        ['인사팀', '인사 관리 업무'],
        ['총무팀', '총무 업무'],
        ['영업팀', '영업 관련 업무'],
        ['안전팀', '안전 관리 업무'],
        ['품질팀', '품질 관리 업무']
    ]
    create_table_with_style(doc, len(dept_data), 2, dept_data)
    
    doc.add_paragraph()
    doc.add_paragraph('등록 가능한 직급:')
    position_data = [
        ['직급명', '권한'],
        ['대표', '최고 관리자'],
        ['상무', '상급 관리자'],
        ['이사', '관리자'],
        ['부장', '부서 관리자'],
        ['차장', '중간 관리자'],
        ['과장', '팀 관리자'],
        ['주임', '담당자'],
        ['사원', '일반 직원']
    ]
    create_table_with_style(doc, len(position_data), 2, position_data)
    
    # 4. 직원 기능
    add_heading_numbered(doc, '4. 직원 기능', level=1)
    
    add_heading_numbered(doc, '4.1 대시보드', level=2)
    doc.add_paragraph('로그인 후 직원 대시보드에서 확인할 수 있는 정보:')
    doc.add_paragraph('• 올해 총 휴가 일수')
    doc.add_paragraph('• 사용한 휴가 일수')
    doc.add_paragraph('• 남은 휴가 일수')
    doc.add_paragraph('• 최근 휴가 신청 내역')
    doc.add_paragraph('• 대기중인 휴가 신청 수')
    
    add_heading_numbered(doc, '4.2 휴가 신청', level=2)
    doc.add_paragraph('휴가 신청 절차:')
    doc.add_paragraph('1. 좌측 메뉴에서 "휴가 신청" 클릭')
    doc.add_paragraph('2. 휴가 정보 입력:')
    doc.add_paragraph('   • 시작일 선택')
    doc.add_paragraph('   • 종료일 선택')
    doc.add_paragraph('   • 휴가 유형 선택 (연차, 반차, 특별휴가)')
    doc.add_paragraph('   • 휴가 사유 입력')
    doc.add_paragraph('3. "신청하기" 버튼 클릭')
    
    doc.add_paragraph()
    doc.add_paragraph('휴가 유형별 특징:')
    vacation_types = [
        ['유형', '설명', '일수 계산'],
        ['연차', '일반 휴가', '실제 근무일 기준'],
        ['반차(오전)', '오전 반일 휴가', '0.5일'],
        ['반차(오후)', '오후 반일 휴가', '0.5일'],
        ['특별휴가', '경조사 등 특별 휴가', '연차 차감 없음']
    ]
    create_table_with_style(doc, len(vacation_types), 3, vacation_types)
    
    add_heading_numbered(doc, '4.3 나의 휴가현황', level=2)
    doc.add_paragraph('나의 휴가현황 페이지에서 제공하는 기능:')
    doc.add_paragraph('• 연도별 휴가 사용 현황 확인')
    doc.add_paragraph('• 휴가 신청 내역 조회')
    doc.add_paragraph('• 기간별 검색 기능')
    doc.add_paragraph('• 상태별 필터링 (전체, 대기중, 승인됨, 반려됨)')
    doc.add_paragraph('• 엑셀 파일 다운로드')
    doc.add_paragraph('• 대기중인 휴가 신청 취소')
    
    add_heading_numbered(doc, '4.4 재직증명서 신청', level=2)
    doc.add_paragraph('재직증명서 신청 방법:')
    doc.add_paragraph('1. 좌측 메뉴에서 "재직증명서 신청" 클릭')
    doc.add_paragraph('2. 신청 정보 입력:')
    doc.add_paragraph('   • 발급 목적 선택')
    doc.add_paragraph('   • 세부 사유 입력')
    doc.add_paragraph('3. "신청하기" 버튼 클릭')
    doc.add_paragraph('4. 관리자 승인 후 PDF 파일 다운로드 가능')
    
    # 5. 관리자 기능
    add_heading_numbered(doc, '5. 관리자 기능', level=1)
    
    add_heading_numbered(doc, '5.1 관리자 대시보드', level=2)
    doc.add_paragraph('관리자 대시보드에서 확인할 수 있는 정보:')
    doc.add_paragraph('• 총 직원 수')
    doc.add_paragraph('• 대기중인 휴가 신청 수')
    doc.add_paragraph('• 대기중인 재직증명서 신청 수')
    doc.add_paragraph('• 이번 달 휴가 사용 현황')
    doc.add_paragraph('• 최근 활동 내역')
    
    add_heading_numbered(doc, '5.2 직원 관리', level=2)
    doc.add_paragraph('직원 관리 기능:')
    doc.add_paragraph('• 전체 직원 목록 조회')
    doc.add_paragraph('• 부서별 직원 필터링')
    doc.add_paragraph('• 직원 정보 수정')
    doc.add_paragraph('• 휴가 일수 조정')
    doc.add_paragraph('• 엑셀 파일로 직원 일괄 등록')
    doc.add_paragraph('• 샘플 파일 다운로드')
    
    add_heading_numbered(doc, '5.3 휴가 관리', level=2)
    doc.add_paragraph('휴가 관리 기능:')
    doc.add_paragraph('• 전체 휴가 신청 내역 조회')
    doc.add_paragraph('• 상태별 필터링 (대기중, 승인됨, 반려됨)')
    doc.add_paragraph('• 부서별 필터링')
    doc.add_paragraph('• 기간별 검색')
    doc.add_paragraph('• 휴가 승인/반려 처리')
    doc.add_paragraph('• 엑셀 파일 다운로드')
    
    doc.add_paragraph()
    doc.add_paragraph('휴가 승인/반려 절차:')
    doc.add_paragraph('1. "처리하기" 버튼 클릭')
    doc.add_paragraph('2. 휴가 정보 확인')
    doc.add_paragraph('3. "승인" 또는 "반려" 선택')
    doc.add_paragraph('4. 반려 시 사유 입력')
    doc.add_paragraph('5. "처리하기" 버튼 클릭')
    
    add_heading_numbered(doc, '5.4 재직증명서 관리', level=2)
    doc.add_paragraph('재직증명서 관리 기능:')
    doc.add_paragraph('• 전체 재직증명서 신청 내역 조회')
    doc.add_paragraph('• 상태별 필터링')
    doc.add_paragraph('• 신청 승인/반려 처리')
    doc.add_paragraph('• PDF 파일 생성 및 다운로드')
    
    add_heading_numbered(doc, '5.5 공휴일 관리', level=2)
    doc.add_paragraph('공휴일 관리 기능:')
    doc.add_paragraph('• 연도별 공휴일 목록 조회')
    doc.add_paragraph('• 신정, 설날, 어린이날, 현충일 등 법정공휴일 자동 등록')
    doc.add_paragraph('• 공휴일 추가/삭제')
    doc.add_paragraph('• 휴가 일수 계산 시 공휴일 자동 제외')
    
    # 6. 고급 기능
    add_heading_numbered(doc, '6. 고급 기능', level=1)
    
    add_heading_numbered(doc, '6.1 기간 검색', level=2)
    doc.add_paragraph('휴가 현황 페이지에서 기간별 검색 방법:')
    doc.add_paragraph('1. 검색 영역에서 시작일 입력')
    doc.add_paragraph('2. 종료일 입력')
    doc.add_paragraph('3. 필요시 상태, 부서, 연도 선택')
    doc.add_paragraph('4. "검색" 버튼 클릭')
    doc.add_paragraph('5. 조건에 맞는 휴가 신청 내역 표시')
    
    add_heading_numbered(doc, '6.2 엑셀 출력', level=2)
    doc.add_paragraph('엑셀 파일 다운로드 기능:')
    doc.add_paragraph('• 검색 조건에 맞는 데이터만 엑셀로 출력')
    doc.add_paragraph('• 파일명에 생성 날짜/시간 자동 포함')
    doc.add_paragraph('• 한글 파일명 지원')
    doc.add_paragraph('• .xlsx 형식으로 저장')
    
    doc.add_paragraph()
    doc.add_paragraph('엑셀 파일에 포함되는 정보:')
    excel_info = [
        '전직원 휴가현황: 이름, 아이디, 부서, 직급, 휴가기간, 일수, 유형, 사유, 상태, 신청일시, 승인일시',
        '개인별 휴가현황: 휴가기간, 일수, 유형, 사유, 상태, 신청일시, 승인일시'
    ]
    for info in excel_info:
        doc.add_paragraph(f'• {info}')
    
    add_heading_numbered(doc, '6.3 계정 복구', level=2)
    doc.add_paragraph('아이디 찾기:')
    doc.add_paragraph('1. 로그인 페이지에서 "아이디 찾기" 클릭')
    doc.add_paragraph('2. 이름과 이메일 입력')
    doc.add_paragraph('3. 일치하는 계정이 있으면 아이디 표시')
    
    doc.add_paragraph()
    doc.add_paragraph('비밀번호 찾기:')
    doc.add_paragraph('1. 로그인 페이지에서 "비밀번호 찾기" 클릭')
    doc.add_paragraph('2. 아이디와 이메일 입력')
    doc.add_paragraph('3. 계정 확인 후 임시 비밀번호 발급')
    doc.add_paragraph('4. 로그인 후 비밀번호 변경 권장')
    
    # 7. 시스템 관리
    add_heading_numbered(doc, '7. 시스템 관리', level=1)
    
    add_heading_numbered(doc, '7.1 데이터베이스', level=2)
    doc.add_paragraph('시스템은 PostgreSQL 데이터베이스를 사용하며, 연결 실패 시 SQLite로 자동 전환됩니다.')
    doc.add_paragraph('• 영구 데이터 저장')
    doc.add_paragraph('• 자동 백업 기능')
    doc.add_paragraph('• 안전한 암호화 저장')
    
    add_heading_numbered(doc, '7.2 보안', level=2)
    doc.add_paragraph('시스템 보안 기능:')
    doc.add_paragraph('• 비밀번호 해시 암호화')
    doc.add_paragraph('• 세션 기반 인증')
    doc.add_paragraph('• CSRF 보호')
    doc.add_paragraph('• 권한 기반 접근 제어')
    
    # 8. 문제해결 가이드
    add_heading_numbered(doc, '8. 문제해결 가이드', level=1)
    
    add_heading_numbered(doc, '8.1 로그인 문제', level=2)
    troubleshoot_login = [
        ['문제', '해결방법'],
        ['아이디/비밀번호를 잊음', '계정 복구 기능 사용'],
        ['로그인이 안됨', '아이디와 비밀번호 정확성 확인'],
        ['페이지가 안 열림', '인터넷 연결 및 브라우저 확인']
    ]
    create_table_with_style(doc, len(troubleshoot_login), 2, troubleshoot_login)
    
    add_heading_numbered(doc, '8.2 휴가 신청 문제', level=2)
    troubleshoot_vacation = [
        ['문제', '해결방법'],
        ['휴가 일수가 부족', '관리자에게 휴가 일수 조정 요청'],
        ['중복 휴가 신청', '기존 휴가와 겹치지 않는 날짜 선택'],
        ['휴가 신청이 안됨', '필수 항목 모두 입력 확인']
    ]
    create_table_with_style(doc, len(troubleshoot_vacation), 2, troubleshoot_vacation)
    
    add_heading_numbered(doc, '8.3 기타 문제', level=2)
    troubleshoot_others = [
        ['문제', '해결방법'],
        ['엑셀 다운로드 안됨', '브라우저 팝업 차단 해제'],
        ['페이지 로딩 느림', '브라우저 캐시 삭제'],
        ['모바일에서 화면 깨짐', '브라우저 새로고침 또는 재시작']
    ]
    create_table_with_style(doc, len(troubleshoot_others), 2, troubleshoot_others)
    
    # 문서 저장
    filename = f'에스에스전력_휴가관리시스템_매뉴얼_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    return filename

if __name__ == '__main__':
    try:
        filename = create_system_manual()
        print(f"✅ 시스템 매뉴얼이 생성되었습니다: {filename}")
        print("📁 파일 위치:", os.path.abspath(filename))
    except Exception as e:
        print(f"❌ 매뉴얼 생성 중 오류 발생: {e}")